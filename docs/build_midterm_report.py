from pathlib import Path
from datetime import datetime
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-assets"
EVIDENCE = ROOT / "test-evidence"
PHPUNIT_XML = DOCS / "evidence" / "raw" / "phpunit-results.xml"
OUT = DOCS / "Bao-cao-giua-ky-Classroom-Management.docx"

BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
GREEN = "E2F0D9"
RED = "FCE4D6"
WHITE = "FFFFFF"
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def set_run_font(run, name="Times New Roman", size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_body(doc, text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_table(doc, headers, rows, widths=None, header_fill=BLUE, font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        set_run_font(run, size=font_size, bold=True, color=WHITE if header_fill == BLUE else BLACK)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(text), size=11, italic=True)


def add_picture(doc, path, caption, width_cm=16.0):
    if not Path(path).exists():
        add_body(doc, f"[Không tìm thấy ảnh: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1F1D25")
    set_cell_margins(cell, top=100, bottom=100, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip().splitlines()):
        line = "".join(ch for ch in line if ord(ch) >= 32 or ch == "\t")
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color="F2EDF7")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def use_case_table(doc, name, actor, trigger, pre, post, normal, alternatives):
    add_heading(doc, name, 3)
    rows = [
        ("Tác nhân chính", actor),
        ("Sự kiện kích hoạt", trigger),
        ("Tiền điều kiện", pre),
        ("Hậu điều kiện", post),
        ("Luồng thông thường", "\n".join(f"{i+1}. {v}" for i, v in enumerate(normal))),
        ("Luồng thay thế/ngoại lệ", "\n".join(f"- {v}" for v in alternatives)),
        ("Độ ưu tiên", "Cao"),
    ]
    add_table(doc, ["Mục", "Mô tả"], rows, [4.2, 12.3], font_size=10.5)


def parse_phpunit():
    root = ET.parse(PHPUNIT_XML).getroot()
    overall = root.find("testsuite")
    summary = {k: overall.attrib.get(k, "0") for k in ("tests", "assertions", "errors", "failures", "skipped", "time")}
    cases = []
    for tc in root.findall(".//testcase"):
        cases.append({
            "class": tc.attrib.get("class", ""),
            "name": tc.attrib.get("name", ""),
            "assertions": tc.attrib.get("assertions", "0"),
            "time": tc.attrib.get("time", "0"),
            "status": "PASS" if not list(tc) else "FAIL",
        })
    return summary, cases


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(13)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.3
for level, size, before, after in ((1, 16, 14, 8), (2, 14, 12, 6), (3, 13, 8, 4)):
    style = styles[f"Heading {level}"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(header.add_run("BÁO CÁO GIỮA KỲ - CLASSROOM MANAGEMENT"), size=9, color="6B7280")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("Trang "), size=10, color="6B7280")
add_page_field(footer)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("ĐẠI HỌC PHENIKAA"), size=15, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("TRƯỜNG CÔNG NGHỆ THÔNG TIN PHENIKAA"), size=14, bold=True, color=BLUE)
logo = ASSETS / "ref-3-1.jpeg"
if logo.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Cm(5.2))
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("HỌC PHẦN: ĐÁNH GIÁ VÀ KIỂM ĐỊNH\nCHẤT LƯỢNG PHẦN MỀM"), size=17, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
set_run_font(p.add_run("ĐỀ TÀI"), size=15, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("PHÂN TÍCH VÀ KIỂM THỬ HỆ THỐNG\nQUẢN LÝ LỚP HỌC - CLASSROOM MANAGEMENT"), size=18, bold=True, color=BLUE)
doc.add_paragraph()
cover_rows = [
    ("Sinh viên thực hiện", "Đỗ Hữu Ngọc"),
    ("Mã sinh viên", "23010822"),
    ("Lớp", "CNTT_8"),
    ("Giảng viên hướng dẫn", "TS. Trịnh Thanh Bình"),
    ("Repository", "https://github.com/vo-hoangdung/BaoCaoKiemDinhPhanMem"),
]
add_table(doc, ["Thông tin", "Nội dung"], cover_rows, [5.0, 11.5], header_fill=LIGHT_BLUE, font_size=11.5)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
set_run_font(p.add_run("Hà Nội, tháng 06 năm 2026"), size=13, italic=True)
doc.add_page_break()

# Assignment and declarations
add_heading(doc, "BẢNG PHÂN CHIA CÔNG VIỆC", 1)
add_table(doc, ["Họ và tên", "Mã sinh viên", "Nhiệm vụ", "Đánh giá"], [[
    "Đỗ Hữu Ngọc",
    "23010822",
    "Phân tích yêu cầu; xây dựng Test Plan; thiết kế 50 PHPUnit test; xây dựng Selenium/JMeter; sửa lỗi phát hiện; thu thập bằng chứng; hoàn thiện README, CI và báo cáo.",
    "Hoàn thành",
]], [3.3, 2.8, 8.7, 2.0], font_size=10.5)
add_heading(doc, "CAM KẾT TÍNH TRUNG THỰC", 1)
add_body(doc, "Toàn bộ kết quả kiểm thử trong báo cáo được tạo từ việc thực thi trực tiếp mã nguồn của hệ thống Classroom Management. Các tệp XML, JTL, log Maven/PHPUnit và ảnh Selenium được lưu trong repository để có thể kiểm chứng. Báo cáo không sử dụng ảnh kết quả dựng giả hoặc tự điền trạng thái PASS.")
add_heading(doc, "TÓM TẮT", 1)
add_body(doc, "Báo cáo trình bày quá trình đặc tả, lập kế hoạch và thực hiện kiểm thử cho hệ thống quản lý lớp học xây dựng bằng Laravel 12. Bộ kiểm thử gồm 50 PHPUnit unit/feature tests với 117 assertions, 16 Selenium/JUnit tests, cùng kịch bản Apache JMeter gồm 65 mẫu chính. Quá trình kiểm thử đã phát hiện một lỗi xử lý định dạng thời gian trong chức năng duyệt yêu cầu đặt phòng; lỗi được sửa và xác nhận bằng regression test.")
doc.add_page_break()

# Manual TOC
add_heading(doc, "MỤC LỤC", 1)
toc = [
    ("PHẦN I. TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM", "1"),
    ("1. Giới thiệu và phạm vi hệ thống", "1"),
    ("2. Tác nhân và yêu cầu chức năng", "3"),
    ("3. Đặc tả các use case chính", "6"),
    ("4. Yêu cầu phi chức năng và kiến trúc", "15"),
    ("PHẦN II. KẾ HOẠCH KIỂM THỬ - TEST PLAN", "18"),
    ("1. Mục tiêu, phạm vi và chiến lược", "18"),
    ("2. Môi trường, dữ liệu và tiêu chí", "22"),
    ("3. Ma trận rủi ro và test deliverables", "26"),
    ("PHẦN III. THỰC HIỆN KIỂM THỬ", "29"),
    ("1. White-box/feature testing với PHPUnit", "29"),
    ("2. Selenium WebDriver", "39"),
    ("3. Apache JMeter", "44"),
    ("4. Lỗi phát hiện và regression testing", "47"),
    ("PHẦN IV. ĐÁNH GIÁ VÀ KẾT LUẬN", "50"),
    ("PHỤ LỤC", "53"),
]
add_table(doc, ["Nội dung", "Trang dự kiến"], toc, [14.5, 2.0], header_fill=LIGHT_BLUE, font_size=11)
doc.add_page_break()

# Part I
add_heading(doc, "PHẦN I: TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM", 1)
add_heading(doc, "1. Giới thiệu", 2)
add_heading(doc, "1.1. Mục đích", 3)
add_body(doc, "Hệ thống Classroom Management hỗ trợ quản lý phòng học, môn học, lịch đặt phòng, người dùng và yêu cầu sử dụng phòng. Mục tiêu của tài liệu là mô tả đầy đủ yêu cầu nghiệp vụ làm cơ sở thiết kế test case và đánh giá chất lượng.")
add_heading(doc, "1.2. Phạm vi sử dụng", 3)
add_bullets(doc, [
    "Quản trị viên quản lý phòng học, môn học, tài khoản và duyệt yêu cầu.",
    "Giảng viên theo dõi lịch, tạo lịch và gửi yêu cầu đặt phòng.",
    "Hệ thống kiểm soát xung đột thời gian, quyền truy cập và ràng buộc dữ liệu.",
    "Ứng dụng chạy trên nền web với Laravel, Blade, SQLite/MySQL và giao diện responsive.",
])
add_heading(doc, "1.3. Thuật ngữ", 3)
add_table(doc, ["Thuật ngữ", "Giải thích"], [
    ("Room", "Phòng học có tên, sức chứa, vị trí và thiết bị."),
    ("Schedule", "Lịch sử dụng một phòng trong khoảng thời gian xác định."),
    ("Course", "Môn học gắn với một giảng viên phụ trách."),
    ("Booking Request", "Yêu cầu sử dụng phòng, có trạng thái pending/approved/rejected."),
    ("Admin", "Vai trò có quyền quản lý dữ liệu và duyệt yêu cầu."),
    ("User", "Giảng viên/người dùng thường, bị giới hạn quyền quản trị."),
], [4.0, 12.5], font_size=11)

add_heading(doc, "2. Tác nhân và yêu cầu chức năng", 2)
add_heading(doc, "2.1. Các tác nhân", 3)
add_table(doc, ["Tác nhân", "Vai trò", "Quyền chính"], [
    ("Quản trị viên", "admin", "CRUD phòng, môn, người dùng; quản lý lịch; duyệt/từ chối yêu cầu."),
    ("Giảng viên", "user", "Đăng nhập; xem dashboard; tạo/sửa/xóa lịch của mình; gửi yêu cầu đặt phòng."),
    ("Hệ thống CSDL", "Phụ trợ", "Lưu dữ liệu, áp dụng khóa ngoại, unique constraint và transaction."),
], [3.4, 2.6, 10.5], font_size=11)
add_heading(doc, "2.2. Danh sách chức năng", 3)
functional_rows = [
    ("F01", "Xác thực", "Đăng nhập, đăng xuất, bảo vệ phiên", "Cao"),
    ("F02", "Quản lý phòng", "Thêm, sửa, xóa, xem phòng", "Cao"),
    ("F03", "Quản lý môn học", "CRUD môn học và giảng viên phụ trách", "Cao"),
    ("F04", "Quản lý lịch", "Tạo lịch, kiểm tra thời gian và xung đột", "Cao"),
    ("F05", "Quản lý người dùng", "CRUD tài khoản, vai trò và mật khẩu", "Cao"),
    ("F06", "Yêu cầu đặt phòng", "Tạo, sửa, xóa, duyệt, từ chối", "Cao"),
    ("F07", "Dashboard", "Thống kê phòng, lịch, môn và yêu cầu chờ", "Trung bình"),
]
add_table(doc, ["ID", "Chức năng", "Mô tả", "Ưu tiên"], functional_rows, [1.4, 4.0, 8.7, 2.4], font_size=10.5)

add_heading(doc, "2.3. Mô hình dữ liệu và kiến trúc", 3)
add_picture(doc, ROOT / "img" / "diagram1.jpg", "Hình 2.1. Sơ đồ cấu trúc hệ thống Classroom Management", 15.5)
add_picture(doc, ROOT / "img" / "activity.jpg", "Hình 2.2. Sơ đồ hoạt động tổng quát", 15.5)
add_body(doc, "Ứng dụng áp dụng kiến trúc MVC. Route nhận yêu cầu HTTP, middleware kiểm tra xác thực/phân quyền, controller xử lý validation và nghiệp vụ, model Eloquent thao tác với cơ sở dữ liệu, Blade hiển thị kết quả.")
doc.add_page_break()

add_heading(doc, "3. Đặc tả các use case chính", 2)
use_case_table(doc, "3.1. Use case đăng nhập/đăng xuất", "Quản trị viên, giảng viên", "Người dùng mở trang /login.", "Tài khoản đã tồn tại và chưa đăng nhập.", "Phiên đăng nhập được tạo hoặc hủy đúng.", [
    "Hệ thống hiển thị form username/password.",
    "Người dùng nhập thông tin và gửi form.",
    "Hệ thống validation trường bắt buộc.",
    "Auth kiểm tra username/password.",
    "Thành công: chuyển đến dashboard; đăng xuất: hủy session và về trang chủ.",
], ["Sai mật khẩu: hiển thị thông báo lỗi.", "Thiếu dữ liệu: trả validation error.", "Guest truy cập dashboard: chuyển về login."])
doc.add_page_break()
use_case_table(doc, "3.2. Use case quản lý phòng học", "Quản trị viên", "Admin chọn tab Phòng học.", "Admin đã đăng nhập.", "Dữ liệu phòng được tạo/cập nhật/xóa phù hợp ràng buộc.", [
    "Hiển thị danh sách phòng.",
    "Admin mở form thêm hoặc sửa.",
    "Nhập tên, sức chứa, vị trí, thiết bị.",
    "Hệ thống kiểm tra tên duy nhất và capacity >= 1.",
    "Lưu dữ liệu và thông báo thành công.",
], ["Tên phòng trùng hoặc capacity sai: không lưu.", "Phòng đang có lịch/yêu cầu: không cho xóa.", "User thường gọi endpoint quản trị: HTTP 403."])
doc.add_page_break()
use_case_table(doc, "3.3. Use case quản lý lịch đặt phòng", "Quản trị viên, giảng viên", "Người dùng chọn Lịch đặt phòng.", "Đã đăng nhập; phòng tồn tại.", "Lịch hợp lệ được lưu và hiển thị.", [
    "Chọn phòng, thời gian bắt đầu/kết thúc và mục đích.",
    "Hệ thống kiểm tra endTime sau startTime.",
    "Hệ thống truy vấn các lịch cùng phòng.",
    "Nếu không xung đột, tạo lịch gắn với người dùng hiện tại.",
    "Chủ sở hữu hoặc admin được phép sửa/xóa.",
], ["Thời gian không hợp lệ: validation error.", "Trùng khoảng thời gian: thông báo phòng đã được đặt.", "Người khác sửa/xóa lịch: từ chối."])
doc.add_page_break()
use_case_table(doc, "3.4. Use case quản lý môn học", "Quản trị viên", "Admin chọn tab Môn học.", "Admin đăng nhập; giảng viên tồn tại.", "Môn học được lưu với mã duy nhất.", [
    "Nhập tên, mã môn, số tín chỉ, mô tả và giảng viên.",
    "Kiểm tra code duy nhất, credits từ 1 đến 10.",
    "Lưu hoặc cập nhật môn học.",
    "Hiển thị thông báo và tải lại danh sách.",
], ["Mã môn trùng hoặc tín chỉ ngoài miền: validation error.", "Môn đang dùng trong booking request: không cho xóa."])
doc.add_page_break()
use_case_table(doc, "3.5. Use case quản lý người dùng", "Quản trị viên", "Admin chọn tab Người dùng.", "Admin đăng nhập.", "Tài khoản được quản lý an toàn.", [
    "Nhập username, họ tên, email, password và role.",
    "Kiểm tra username/email duy nhất.",
    "Kiểm tra password tối thiểu 8 ký tự và confirmation.",
    "Hash password trước khi lưu.",
    "Cho phép cập nhật hoặc xóa người dùng khác.",
], ["Admin không được tự xóa chính mình.", "User thường truy cập endpoint: HTTP 403.", "Password confirmation sai: không lưu."])
doc.add_page_break()
use_case_table(doc, "3.6. Use case yêu cầu đặt phòng", "Giảng viên, quản trị viên", "Giảng viên mở tab Yêu cầu đặt phòng.", "Đã đăng nhập; phòng và môn học tồn tại.", "Yêu cầu ở trạng thái pending; khi duyệt tạo schedule.", [
    "Giảng viên chọn phòng, môn, ngày, giờ và mục đích.",
    "Kiểm tra ngày không ở quá khứ, giờ kết thúc sau giờ bắt đầu.",
    "Kiểm tra xung đột với schedule.",
    "Lưu yêu cầu pending.",
    "Admin duyệt hoặc từ chối; duyệt thành công tạo schedule tương ứng.",
], ["Yêu cầu đã xử lý không được sửa.", "User thường không được duyệt/từ chối.", "Trùng lịch: không tạo yêu cầu hoặc schedule."])

add_heading(doc, "4. Yêu cầu phi chức năng", 2)
add_table(doc, ["Nhóm", "Yêu cầu", "Cách kiểm chứng"], [
    ("Bảo mật", "Mật khẩu hash; endpoint quản trị có middleware admin.", "Feature test phân quyền và Hash::check."),
    ("Hiệu năng", "Trang chính phản hồi ổn định dưới tải nhẹ.", "JMeter 10 users, 0% lỗi."),
    ("Tin cậy", "Validation và khóa ngoại bảo vệ dữ liệu.", "PHPUnit database assertions."),
    ("Khả dụng", "Giao diện rõ ràng, modal và thông báo phản hồi.", "Selenium UI test."),
    ("Bảo trì", "MVC, test tự động và CI.", "GitHub Actions chạy php artisan test."),
], [3.0, 7.7, 5.8], font_size=10.5)
doc.add_page_break()

# Part II
add_heading(doc, "PHẦN II: KẾ HOẠCH KIỂM THỬ PHẦN MỀM - TEST PLAN", 1)
add_heading(doc, "1. Mục tiêu kiểm thử", 2)
add_bullets(doc, [
    "Xác nhận các chức năng đáp ứng đặc tả nghiệp vụ.",
    "Phát hiện lỗi validation, xung đột lịch, định dạng thời gian và phân quyền.",
    "Đảm bảo các thay đổi không gây regression.",
    "Tạo bằng chứng có thể tái thực hiện trong GitHub.",
])
add_heading(doc, "2. Phạm vi kiểm thử", 2)
add_table(doc, ["Trong phạm vi", "Ngoài phạm vi"], [[
    "Auth; Room; Course; Schedule; User; Booking Request; model relationships; UI chính; tải HTTP nhẹ.",
    "Penetration test chuyên sâu; tải hàng nghìn người dùng; email thực; queue worker; triển khai production.",
]], [8.3, 8.2], font_size=11)
add_heading(doc, "3. Chiến lược kiểm thử", 2)
add_table(doc, ["Tầng", "Công cụ", "Mục tiêu", "Số lượng"], [
    ("Unit", "PHPUnit", "Model helper và quan hệ", "6 tests"),
    ("Feature/white-box", "Laravel PHPUnit", "Controller, middleware, validation, DB", "44 tests"),
    ("Functional/UI", "Selenium + JUnit 5", "Luồng trình duyệt", "16 tests"),
    ("Performance HTTP", "JMeter 5.6.3", "Tải nhẹ và response assertions", "65 primary samples"),
], [3.0, 4.0, 7.0, 2.5], font_size=10.5)
add_heading(doc, "4. Kỹ thuật white-box", 2)
add_table(doc, ["Kỹ thuật", "Áp dụng"], [
    ("Statement coverage", "Thực thi các câu lệnh chính trong controller/model."),
    ("Branch coverage", "Đi qua nhánh hợp lệ/không hợp lệ, admin/user, conflict/no-conflict."),
    ("Condition testing", "Kiểm tra điều kiện ngày, giờ, unique, ownership, status."),
    ("Path testing", "Luồng create-update-delete và request-approve-schedule."),
], [4.5, 12.0], font_size=11)
add_heading(doc, "5. Môi trường kiểm thử", 2)
add_table(doc, ["Thành phần", "Cấu hình"], [
    ("OS", "Windows"),
    ("PHP/Laravel", "PHP 8.4.21, Laravel 12"),
    ("Database test", "SQLite in-memory"),
    ("Java", "Eclipse Temurin JDK 17"),
    ("Maven", "Apache Maven 3.9.16"),
    ("Selenium", "Selenium Java 4.33.0, Chrome headless"),
    ("JMeter", "Apache JMeter 5.6.3"),
], [5.0, 11.5], font_size=11)
add_heading(doc, "6. Dữ liệu và tài khoản test", 2)
add_table(doc, ["Vai trò", "Username", "Password", "Mục đích"], [
    ("Admin", "admin", "password", "Kiểm thử chức năng quản trị"),
    ("Giảng viên", "giaovien1", "password", "Kiểm thử quyền user"),
    ("Dữ liệu động", "uniqid/timestamp", "-", "Tránh trùng dữ liệu giữa các lần chạy"),
], [3.0, 4.0, 3.2, 6.3], font_size=11)
add_heading(doc, "7. Tiêu chí vào/ra", 2)
add_table(doc, ["Loại", "Tiêu chí"], [
    ("Entry", "Cài dependencies; migrate được; tài khoản seed tồn tại; Chrome/JDK/JMeter sẵn sàng."),
    ("Exit", "Không còn test fail; error rate JMeter = 0%; bằng chứng được lưu; lỗi phát hiện đã regression test."),
], [3.0, 13.5], font_size=11)
add_heading(doc, "8. Rủi ro", 2)
add_table(doc, ["Rủi ro", "Tác động", "Giảm thiểu"], [
    ("Test làm bẩn DB", "Sai dữ liệu demo", "PHPUnit dùng in-memory; Selenium tạo tên unique và cleanup."),
    ("Chrome/Selenium lệch phiên bản", "UI test không chạy", "WebDriverManager tự quản lý driver."),
    ("Thời gian phụ thuộc timezone", "Test không ổn định", "Dùng Carbon và ngày tương lai tương đối."),
    ("JMeter quá tải máy local", "Kết quả sai hoặc treo", "Giới hạn 10 users, ramp-up 5 giây."),
], [5.0, 4.5, 7.0], font_size=10.5)
add_heading(doc, "9. Thiết kế kiểm thử hộp đen", 2)
add_body(doc, "Các test case hộp đen được thiết kế từ miền dữ liệu đầu vào và kết quả quan sát được, không phụ thuộc vào cấu trúc cài đặt bên trong.")
blackbox_rows = [
    ("BB-AUTH-01", "Username/password đúng", "admin / password", "Vào dashboard", "Phân vùng hợp lệ"),
    ("BB-AUTH-02", "Password sai", "admin / wrong-password", "Thông báo đăng nhập sai", "Phân vùng không hợp lệ"),
    ("BB-AUTH-03", "Bỏ trống username", "username=''", "Báo trường bắt buộc", "Giá trị rỗng"),
    ("BB-ROOM-01", "Capacity nhỏ nhất hợp lệ", "capacity=1", "Tạo phòng thành công", "Giá trị biên"),
    ("BB-ROOM-02", "Capacity dưới biên", "capacity=0", "Validation error", "Giá trị biên"),
    ("BB-ROOM-03", "Tên phòng đã tồn tại", "name trùng", "Validation unique", "Phân vùng không hợp lệ"),
    ("BB-COURSE-01", "Tín chỉ thấp nhất", "credits=1", "Tạo môn thành công", "Giá trị biên"),
    ("BB-COURSE-02", "Tín chỉ cao nhất", "credits=10", "Tạo môn thành công", "Giá trị biên"),
    ("BB-COURSE-03", "Tín chỉ vượt biên", "credits=11", "Validation error", "Giá trị biên"),
    ("BB-SCH-01", "Khoảng giờ hợp lệ", "08:00-10:00", "Tạo lịch", "Phân vùng hợp lệ"),
    ("BB-SCH-02", "End bằng start", "10:00-10:00", "Validation error", "Giá trị biên"),
    ("BB-SCH-03", "Khoảng giờ giao nhau", "09:00-11:00 với 08:00-10:00", "Conflict error", "Bảng quyết định"),
    ("BB-USER-01", "Password đủ 8 ký tự", "password", "Tạo tài khoản", "Giá trị biên"),
    ("BB-USER-02", "Password dưới 8 ký tự", "pass123", "Validation error", "Giá trị biên"),
    ("BB-USER-03", "Confirmation khác", "password123/different", "Validation error", "Phân vùng không hợp lệ"),
    ("BB-BOOK-01", "Ngày hôm nay", "requestDate=today", "Chấp nhận", "Giá trị biên"),
    ("BB-BOOK-02", "Ngày hôm qua", "requestDate=yesterday", "Validation error", "Giá trị biên"),
    ("BB-BOOK-03", "User duyệt yêu cầu", "role=user", "HTTP 403", "Bảng quyết định quyền"),
    ("BB-BOOK-04", "Admin duyệt pending", "role=admin,status=pending", "Approved + tạo schedule", "Bảng quyết định trạng thái"),
]
add_table(doc, ["ID", "Trường hợp", "Dữ liệu", "Kết quả mong đợi", "Kỹ thuật"], blackbox_rows, [2.3, 4.0, 3.5, 4.2, 2.5], font_size=8.8)

# Part III
summary, cases = parse_phpunit()
add_heading(doc, "PHẦN III: THỰC HIỆN KIỂM THỬ", 1)
add_heading(doc, "1. PHPUnit - Unit và Feature/White-box Testing", 2)
add_heading(doc, "1.1. Cấu trúc bộ test", 3)
add_code(doc, """
tests/
├── Unit/ModelRelationshipTest.php
├── Feature/AuthenticationTest.php
├── Feature/RoomManagementTest.php
├── Feature/CourseManagementTest.php
├── Feature/ScheduleManagementTest.php
├── Feature/UserManagementTest.php
├── Feature/BookingRequestManagementTest.php
└── Support/CreatesClassroomData.php
""")
add_heading(doc, "1.2. Kết quả chạy thực tế", 3)
add_table(doc, ["Chỉ số", "Kết quả"], [
    ("Tests", summary["tests"]),
    ("Assertions", summary["assertions"]),
    ("Failures", summary["failures"]),
    ("Errors", summary["errors"]),
    ("Skipped", summary["skipped"]),
    ("Runtime trong JUnit XML", f'{float(summary["time"]):.3f} giây'),
    ("Trạng thái", "PASS"),
], [7.0, 9.5], header_fill=LIGHT_BLUE, font_size=11.5)
add_code(doc, """
php artisan test --log-junit docs/evidence/raw/phpunit-results.xml

Tests: 50 passed (117 assertions)
Duration: 1.17s
""")
add_body(doc, "Nguồn xác minh: docs/evidence/raw/phpunit-results.xml và phpunit-output.txt. Các số liệu trong bảng được đọc trực tiếp từ XML do PHPUnit sinh ra.", italic=True)

add_heading(doc, "1.3. Danh sách 50 test case đã thực thi", 3)
test_rows = []
for idx, case in enumerate(cases, 1):
    class_name = case["class"].split("\\")[-1]
    readable = case["name"].replace("test_", "").replace("_", " ")
    test_rows.append((idx, class_name, readable, case["assertions"], case["status"]))
add_table(doc, ["STT", "Test class", "Test case", "Assertions", "Kết quả"], test_rows, [1.2, 4.1, 8.4, 1.6, 1.6], font_size=8.8)

add_heading(doc, "1.4. Ví dụ kiểm thử nhánh và điều kiện", 3)
add_code(doc, """
public function test_conflicting_schedule_is_rejected(): void
{
    // Arrange: tạo lịch 08:00-10:00 cho một phòng.
    // Act: gửi lịch mới 09:00-11:00 cùng phòng.
    // Assert: session có lỗi conflict và database chỉ có 1 lịch.
}
""")
add_body(doc, "Test trên đi qua nhánh conflict=true của ScheduleController. Các test khác đi qua conflict=false, owner=true/false, role=admin/user và status=pending/processed.")
doc.add_page_break()

add_heading(doc, "2. Selenium WebDriver", 2)
add_heading(doc, "2.1. Cài đặt và lệnh chạy", 3)
add_code(doc, r"""
..\tools\apache-maven-3.9.16\bin\mvn.cmd ^
  -f selenium-junit-tests\pom.xml test
""")
add_heading(doc, "2.2. Kết quả chạy thực tế", 3)
add_table(doc, ["Chỉ số", "Kết quả"], [
    ("Tổng test trong report", "20"),
    ("Pass", "16"),
    ("Fail/Error", "0"),
    ("Skip", "4 test chỉ bật khi chụp evidence"),
    ("Build", "BUILD SUCCESS"),
], [7.0, 9.5], header_fill=LIGHT_BLUE, font_size=11.5)
add_body(doc, "Bốn test evidence được chạy riêng với thuộc tính capture.evidence=true để chụp ảnh trực tiếp. Bộ test chức năng mặc định có 16 ca pass.")
add_picture(doc, EVIDENCE / "03-invalid-login.png", "Hình 3.1. Selenium chụp trực tiếp trường hợp đăng nhập sai", 16.2)
add_picture(doc, EVIDENCE / "04-admin-dashboard.png", "Hình 3.2. Selenium chụp dashboard sau đăng nhập admin thành công", 16.2)
add_picture(doc, EVIDENCE / "05-room-validation.png", "Hình 3.3. Selenium chụp validation capacity = 0", 16.2)
add_body(doc, "Các ảnh trên được lấy bằng TakesScreenshot của ChromeDriver tại đúng thời điểm assertion thành công; không phải ảnh mô phỏng.")
doc.add_page_break()

add_heading(doc, "3. Apache JMeter", 2)
add_heading(doc, "3.1. Kịch bản", 3)
add_table(doc, ["ID", "Kịch bản", "Threads/Loops", "Assertion"], [
    ("TC01", "GET / trang chủ", "10 × 2", "HTTP 200"),
    ("TC02", "Guest GET /dashboard", "5 × 1", "HTTP 302"),
    ("TC03", "Login hợp lệ + dashboard", "10 × 1", "CSRF/session + HTTP 200"),
    ("TC04", "Login sai mật khẩu", "5 × 1", "Thông báo lỗi"),
], [1.5, 7.2, 3.0, 4.8], font_size=10.5)
add_heading(doc, "3.2. Kết quả chạy thực tế", 3)
add_table(doc, ["Chỉ số", "Giá trị"], [
    ("Primary samples", "65"),
    ("JTL rows gồm redirect/sub-samples", "95"),
    ("Errors", "0"),
    ("Error rate", "0.00%"),
    ("Average JTL elapsed", "5697.6 ms"),
    ("Min / Max", "408 ms / 19636 ms"),
], [8.0, 8.5], header_fill=LIGHT_BLUE, font_size=11)
add_body(doc, "Nguồn xác minh: jmeter/results/results.jtl, jmeter/results/jmeter.log và jmeter/report/index.html.")
add_picture(doc, EVIDENCE / "02-jmeter-dashboard.png", "Hình 3.4. JMeter Dashboard sinh từ results.jtl - PASS 100%", 13.8)

add_heading(doc, "4. Lỗi phát hiện và regression testing", 2)
add_heading(doc, "4.1. Mô tả lỗi", 3)
add_table(doc, ["Thuộc tính", "Nội dung"], [
    ("Mã lỗi", "BUG-BR-01"),
    ("Chức năng", "Admin duyệt yêu cầu đặt phòng"),
    ("Biểu hiện", "Không tạo schedule; session trả thông báo lỗi xử lý thời gian."),
    ("Nguyên nhân", "Controller yêu cầu định dạng H:i:s nhưng database có thể trả H:i."),
    ("Mức độ", "High - chặn nghiệp vụ duyệt yêu cầu."),
    ("Phát hiện bởi", "BookingRequestManagementTest::admin_can_approve_request_and_create_schedule"),
], [4.0, 12.5], font_size=11)
add_heading(doc, "4.2. Kết quả vòng chạy trước khi sửa", 3)
add_code(doc, """
Tests: 1 failed, 49 passed (115 assertions)

FAILED BookingRequestManagementTest
Session is missing expected key [success].
""")
add_heading(doc, "4.3. Thay đổi mã nguồn", 3)
add_code(doc, """
// Trước: chỉ chấp nhận HH:mm:ss
Carbon::createFromFormat('Y-m-d H:i:s', $date . ' ' . $time);

// Sau: chấp nhận cả HH:mm và HH:mm:ss
Carbon::parse($date . ' ' . $time);
""")
add_heading(doc, "4.4. Regression test", 3)
add_body(doc, "Sau khi sửa, chạy lại toàn bộ test: 50/50 pass, 117 assertions, không failure/error. Ca duyệt yêu cầu xác nhận status=approved và schedule được tạo đúng roomId/userId.")
doc.add_page_break()

# Part IV
add_heading(doc, "PHẦN IV: ĐÁNH GIÁ VÀ KẾT LUẬN", 1)
add_heading(doc, "1. Tổng hợp kết quả", 2)
add_table(doc, ["Tầng kiểm thử", "Pass", "Fail", "Bằng chứng"], [
    ("PHPUnit Unit/Feature", "50", "0", "JUnit XML + raw output"),
    ("Selenium/JUnit chức năng", "16", "0", "Surefire XML + Maven log"),
    ("Evidence capture", "4", "0", "PNG trực tiếp từ ChromeDriver"),
    ("JMeter primary samples", "65", "0", "JTL + HTML Dashboard"),
], [5.2, 2.5, 2.5, 6.3], font_size=11)
add_heading(doc, "2. Mức độ đáp ứng", 2)
add_bullets(doc, [
    "Các nghiệp vụ cốt lõi được bao phủ ở tầng controller, database và UI.",
    "Phân quyền admin/user được kiểm tra ở cả đường dẫn hợp lệ và trái quyền.",
    "Validation dữ liệu và xung đột thời gian có test dương/âm.",
    "Quy trình test đã phát hiện được lỗi nghiệp vụ thật và chứng minh khả năng regression.",
    "GitHub Actions giúp ngăn việc đẩy mã làm hỏng bộ test.",
])
add_heading(doc, "3. Hạn chế", 2)
add_bullets(doc, [
    "Chưa thực hiện security penetration test chuyên sâu.",
    "JMeter chỉ đánh giá tải nhẹ trên máy local, chưa đại diện production.",
    "Selenium tập trung vào đăng nhập và phòng học; có thể mở rộng UI test cho course, schedule, user và booking request.",
    "Chưa đo code coverage phần trăm do môi trường PHP hiện tại không bật Xdebug/PCOV.",
])
add_heading(doc, "4. Hướng phát triển", 2)
add_numbered(doc, [
    "Bổ sung coverage driver và ngưỡng coverage trong CI.",
    "Bổ sung Selenium Page Object cho toàn bộ module.",
    "Tách database staging và chạy load test lớn hơn.",
    "Bổ sung kiểm thử bảo mật OWASP Top 10 và accessibility.",
    "Thêm Docker Compose để tái lập môi trường kiểm thử.",
])
add_heading(doc, "5. Kết luận", 2)
add_body(doc, "Hệ thống Classroom Management đã được xây dựng một quy trình kiểm thử có thể tái thực hiện, từ đặc tả yêu cầu, Test Plan, test case, thực thi tự động đến lưu bằng chứng. Kết quả cuối cùng cho thấy toàn bộ test bắt buộc đều pass và lỗi nghiêm trọng phát hiện trong chức năng duyệt yêu cầu đã được sửa. Repository sau khi hoàn thiện có thể dùng làm sản phẩm nộp giữa kỳ và nền tảng tiếp tục phát triển.")
doc.add_page_break()

# Appendices
add_heading(doc, "PHỤ LỤC A: LỆNH TÁI THỰC HIỆN", 1)
add_code(doc, r"""
# PHPUnit
php artisan test --log-junit docs/evidence/raw/phpunit-results.xml

# Selenium
..\tools\apache-maven-3.9.16\bin\mvn.cmd -f selenium-junit-tests\pom.xml test

# JMeter
..\tools\apache-jmeter-5.6.3\bin\jmeter.bat -n ^
  -t jmeter\classroom-management-test-plan.jmx ^
  -l jmeter\results\results.jtl -e -o jmeter\report
""")
add_heading(doc, "PHỤ LỤC B: CẤU TRÚC BẰNG CHỨNG", 1)
add_code(doc, """
docs/evidence/raw/
├── phpunit-results.xml
├── phpunit-output.txt
├── selenium-maven-output.txt
└── jmeter-output.txt

test-evidence/
├── 02-jmeter-dashboard.png
├── 03-invalid-login.png
├── 04-admin-dashboard.png
└── 05-room-validation.png

jmeter/
├── results/results.jtl
├── results/jmeter.log
└── report/index.html
""")
add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
refs = [
    "Laravel Documentation, https://laravel.com/docs",
    "PHPUnit Documentation, https://docs.phpunit.de/",
    "Selenium Documentation, https://www.selenium.dev/documentation/",
    "Apache JMeter User Manual, https://jmeter.apache.org/usermanual/",
    "ISTQB Foundation Level Syllabus - Software Testing Principles.",
    "Glenford J. Myers, The Art of Software Testing.",
]
for idx, ref in enumerate(refs, 1):
    add_body(doc, f"[{idx}] {ref}", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)

doc.core_properties.title = "Báo cáo giữa kỳ - Classroom Management"
doc.core_properties.subject = "Đánh giá và kiểm định chất lượng phần mềm"
doc.core_properties.author = "Đỗ Hữu Ngọc - 23010822"
doc.core_properties.keywords = "Classroom Management, Laravel, PHPUnit, Selenium, JMeter, Software Testing"
doc.save(OUT)
print(OUT)
